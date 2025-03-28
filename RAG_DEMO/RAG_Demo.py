# Databricks notebook source
# MAGIC %pip install --upgrade --quiet databricks-sdk langchain langchain-community databricks-langchain langchain_core langchain_community langgraph databricks-agents mlflow mlflow-skinny googlemaps pypdf unstructured databricks-vectorsearch python-docx openpyxl googlemaps

# COMMAND ----------

# DBTITLE 1,有効にする為にpythonをリスタート
dbutils.library.restartPython()

# COMMAND ----------

# DBTITLE 1,データの保存場所を指定し、カタログとスキーマを作成

import os
from databricks.sdk.core import DatabricksError
from databricks.sdk import WorkspaceClient
from databricks.sdk.service.vectorsearch import EndpointStatusState, EndpointType
from databricks.sdk.service.serving import EndpointCoreConfigInput, EndpointStateReady
from databricks.sdk.errors import ResourceDoesNotExist, NotFound, PermissionDenied

CURRENT_FOLDER = os.getcwd()
w = WorkspaceClient()

# カタログとスキーマ
UC_CATALOG = 'hhhd_demo_itec'
UC_SCHEMA = 'allowance_payment_rules'

# モデル
UC_MODEL_NAME = f"{UC_CATALOG}.{UC_SCHEMA}.payment_rules_model"

# search endpoint
VECTOR_SEARCH_ENDPOINT = 'payment_rules_vector_search'

# カタログを作成
try:
    _ = w.catalogs.get(UC_CATALOG)
    print(f"PASS: UC catalog `{UC_CATALOG}` exists")
except NotFound as e:
    print(f"`{UC_CATALOG}` does not exist, trying to create...")
    try:
        _ = w.catalogs.create(name=UC_CATALOG)
    except PermissionDenied as e:
        print(f"FAIL: `{UC_CATALOG}` does not exist, and no permissions to create.  Please provide an existing UC Catalog.")
        raise ValueError(f"Unity Catalog `{UC_CATALOG}` does not exist.")
        
# スキーマを作成
try:
    _ = w.schemas.get(full_name=f"{UC_CATALOG}.{UC_SCHEMA}")
    print(f"PASS: UC schema `{UC_CATALOG}.{UC_SCHEMA}` exists")
except NotFound as e:
    print(f"`{UC_CATALOG}.{UC_SCHEMA}` does not exist, trying to create...")
    try:
        _ = w.schemas.create(name=UC_SCHEMA, catalog_name=UC_CATALOG)
        print(f"PASS: UC schema `{UC_CATALOG}.{UC_SCHEMA}` created")
    except PermissionDenied as e:
        print(f"FAIL: `{UC_CATALOG}.{UC_SCHEMA}` does not exist, and no permissions to create.  Please provide an existing UC Schema.")
        raise ValueError("Unity Catalog Schema `{UC_CATALOG}.{UC_SCHEMA}` does not exist.")

# COMMAND ----------

# DBTITLE 1,PDFを読み込んでテキストに変換
from pypdf import PdfReader

pdf_path_1 = "株式会社サンプル 通勤手当支給規程.pdf"
pdf_path_2 = "株式会社サンプル 通勤手当支給規程（全角数字）.pdf"
pdf_path_3 = "株式会社サンプル 通勤手当支給規程（小数点なし）.pdf"
pdf_path_4 = "株式会社サンプル 通勤手当支給規程（漢字数字）.pdf"
pdf_path_5 = "通勤手当支給規程.pdf"

# PDFを読み込んでテキストに変換
with open(pdf_path_5, "rb") as f:
    reader = PdfReader(f)
    text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

# COMMAND ----------

# DBTITLE 1,マルチテキストに変換
from pypdf import PdfReader
from docx import Document
import pandas as pd
import os

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""

    # 通常のテキスト（段落）
    for para in doc.paragraphs:
        if para.text.strip():  # 空白行を削除
            text += para.text.strip() + "\n"

    # テーブルデータの取得
    for table in doc.tables:
        processed_cells = set()  # 既に処理したセルを保存

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if (cell._tc not in processed_cells) and cell_text:  # 結合セルを除外
                    row_data.append(cell_text)
                    processed_cells.add(cell._tc)  # 結合セルを処理済みとして登録
                else:
                    row_data.append("")  # 空白を保持してフォーマットを崩さない

            text += " | ".join(row_data) + "\n"  # テーブルを整形

    return text

def extract_text_from_xlsx(xlsx_path):
    df = pd.read_excel(xlsx_path, engine='openpyxl', sheet_name=None) 
    text = ""
    for sheet_name, sheet_df in df.items():
        text += f"### {sheet_name} シート ###\n"
        text += sheet_df.to_string(index=False, header=True) + "\n"
    return text

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="shift_jis") as f:
        text = f.read()
    return text

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".xls", ".xlsx"]:
        return extract_text_from_xlsx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"対応していないファイル形式です: {ext}")

# file_path_1 = "株式会社サンプル 通勤手当支給規程.pdf"
# file_path_2 = "株式会社サンプル 通勤手当支給規程.docx"
# file_path_3 = "株式会社サンプル 通勤手当支給規程.xlsx"
# file_path_4 = "株式会社サンプル 通勤手当支給規程.txt"   

# text = extract_text(file_path_1)


# COMMAND ----------

# DBTITLE 1,ベクトル検索エンドポイントを作成
# エンドポイントが存在しない場合は作成する
vector_search_endpoints = w.vector_search_endpoints.list_endpoints()
if sum([VECTOR_SEARCH_ENDPOINT == ve.name for ve in vector_search_endpoints]) == 0:
    print(f"Please wait, creating Vector Search endpoint `{VECTOR_SEARCH_ENDPOINT}`.  This can take up to 20 minutes...")
    w.vector_search_endpoints.create_endpoint_and_wait(VECTOR_SEARCH_ENDPOINT, endpoint_type=EndpointType.STANDARD)

# Make sure vector search endpoint is online and ready.
w.vector_search_endpoints.wait_get_endpoint_vector_search_endpoint_online(VECTOR_SEARCH_ENDPOINT)

print(f"PASS: Vector Search endpoint `{VECTOR_SEARCH_ENDPOINT}` exists")

# COMMAND ----------

# DBTITLE 1,エンドポイント一覧表示（必須ではない）
from databricks.sdk import WorkspaceClient

w = WorkspaceClient()

# すべてのベクトル検索エンドポイントをリスト化
vector_search_endpoints = w.vector_search_endpoints.list_endpoints()

if vector_search_endpoints:
    # 取得したエンドポイント一覧を表示
    for endpoint in vector_search_endpoints:
        print(f"Name: {endpoint.name}")
else:
    print("No vector search endpoints found.")

# COMMAND ----------

# DBTITLE 1,Delta Table にデータを保存、インデックス作成 & 同期

from pyspark.sql import SparkSession
from pyspark.sql.functions import monotonically_increasing_id
from databricks.vector_search.client import VectorSearchClient
import re

# SparkSessionを取得
spark = SparkSession.builder.getOrCreate()

workspace_url = SparkSession.getActiveSession().conf.get(
    "spark.databricks.workspaceUrl", None
)

################################################################################################################################

# PDFから取得したテキストをDataFrameに変換
# 行分割
# pdf_texts = [
#     {"chunk_id": i, "chunked_text": text_chunk}
#     for i, text_chunk in enumerate(text.split("\n"))
# ]

################################################################################################################################
# # 「。」分割
# def chunk_text(text):
#     """
#     「。」ごとに文を区切り、リストとして返す。
#     - 文が途中で切れないようにする。
#     - 空の文はリストに含めない。
#     """
#     sentences = re.split(r'(?<=。)\s*', text)  # 「。」の後ろで分割
#     chunks = [s.strip() for s in sentences if s.strip()]  # 空の要素を削除
#     return chunks

# # PDFのテキストを適切な長さのチャンクに変換
# chunked_texts = chunk_text(text)

# # DataFrameに変換
# pdf_texts = [{"chunk_id": i, "chunked_text": chunk} for i, chunk in enumerate(chunked_texts)]

################################################################################################################################
# # 「条」分割
# def extract_sections(text):
#     """
#     文書を適切なセクションごとに分割する。
#     - 文書のタイトル（最初の1行）を取得
#     - 各「第X条 (タイトル)」に基づいてセクションを分割
#     - 「附則」も正しく分割
#     """
#     sections = []

#     # 文書のタイトルを取得（最初の1行）
#     title_match = re.search(r'^(.*?)\n', text)
#     if title_match:
#         sections.append({"title": "文書タイトル", "content": title_match.group(1).strip()})
    
#     # 「第X条 (タイトル)」または「附則」をセクションとして識別
#     matches = re.split(r'(第\s*\d+\s*条.*?|\b附則\b)', text)

#     for i in range(1, len(matches), 2):  # 2つずつ処理
#         title = matches[i].strip()
#         content = matches[i+1].strip() if i+1 < len(matches) else ""

#         # 空のセクションを除外
#         if title and content:
#             sections.append({"title": title, "content": content})

#     return sections

# def chunk_text(text):
#     """
#     条文ごとに適切にチャンクを作成する関数。
#     - 「第〇条」や「第〇項」を検出し、新しいチャンクを作成。
#     - 「。」で適度に分割し、最大3文を1チャンクにする。
#     - 空の文はリストに含めない。
#     """
#     # 改行を統一（余計な改行を削除）
#     text = re.sub(r'\n+', '\n', text).strip()

#     # 「第〇条 (〜)」または「附則」で分割
#     sections = extract_sections(text)

#     chunks = []
#     current_chunk = ""

#     for sec in sections:
#         temp_chunk = sec['title'] + ":" + sec['content']
#         chunks.append(temp_chunk)
#         print(f"title: {sec['title']}\ncontent: {sec['content']}\n")

#     return chunks

# # PDFのテキストを適切な長さのチャンクに変換
# chunked_texts = chunk_text(text)

# # DataFrameに変換
# pdf_texts = [{"chunk_id": i, "chunked_text": chunk} for i, chunk in enumerate(chunked_texts)]

################################################################################################################################

def extract_sections(text):
    """
    文書を適切なセクションごとに分割する。
    - 空白行を境目として各セクションを分割
    - 各「第X条 (タイトル)」や「附則」も個別のエントリとして処理
    """
    sections = []

    # 空白行を境目にして分割
    raw_sections = re.split(r'\n\s*\n+', text.strip())

    for section in raw_sections:
        # 「第X条」や「附則」がタイトルであるか判定
        title_match = re.match(r'^(第\s*\d+\s*条.*?)$', section, re.MULTILINE)

        if title_match:
            title = title_match.group(1).strip()
            content = section[len(title):].strip()  # タイトル以外の本文
            sections.append({"title": title, "content": content})
        else:
            # タイトルがないセクションもそのまま格納
            sections.append({"title": "", "content": section.strip()})

    return sections

def chunk_text(text):
    """
    条文ごとに適切にチャンクを作成する関数。
    - 「第〇条」や「第〇項」を検出し、新しいチャンクを作成
    - 各エントリを独立したチャンクとして保存
    """
    text = re.sub(r'\n+', '\n', text).strip()

    # 空白行で分割
    sections = extract_sections(text)

    chunks = []
    for sec in sections:
        if sec['content']:
            if sec['title']:
                temp_chunk = f"{sec['title']}:\n{sec['content']}"
            else:
                temp_chunk = sec['content']
        else:
            sec['title']
        chunks.append(temp_chunk)

    return chunks

# PDFのテキストを適切な長さのチャンクに変換
chunked_texts = chunk_text(text)

# DataFrameに変換
pdf_texts = [{"chunk_id": i, "chunked_text": chunk} for i, chunk in enumerate(chunked_texts)]

# Deltaテーブルとして保存
DELTA_TABLE_NAME = f"{UC_CATALOG}.{UC_SCHEMA}.payment_rules_chunked_model_bge"

# データをベクトル化
df = spark.createDataFrame(pdf_texts)
df.write.format("delta").mode("overwrite").saveAsTable(DELTA_TABLE_NAME)

spark.sql(
    f"ALTER TABLE {DELTA_TABLE_NAME} SET TBLPROPERTIES (delta.enableChangeDataFeed = true)"
)

# Databricks UI 上で Delta Table (databricks_docs_chunked) を確認できるリンクを表示
# コードを実行したら、出力される URL をクリックし、Databricks UI に飛んで、Delta Table が正しく作成されたことを確認
print(
    f"View Delta Table at: https://{workspace_url}/explore/data/{UC_CATALOG}/{UC_SCHEMA}/{DELTA_TABLE_NAME.split('.')[-1]}"
)

# ベクトル検索インデックス
CHUNKS_VECTOR_INDEX = f"{UC_CATALOG}.{UC_SCHEMA}.payment_rules_chunked_index_new_model_bge"

# ベクトル検索クライアントを取得
vsc = VectorSearchClient()

# ベクトル検索インデックス（CHUNKS_VECTOR_INDEX）の作成を開始
# コード実行後、URL をクリックして Databricks UI でインデックスが作成されているか確認
print(
    f"Embedding docs & creating Vector Search Index, this will take ~5 - 10 minutes.\nView Index Status at: https://{workspace_url}/explore/data/{UC_CATALOG}/{UC_SCHEMA}/{CHUNKS_VECTOR_INDEX.split('.')[-1]}"
)

# インデックスが存在している場合は作成しない
try:
    # インデックス作成 & 同期
    index = vsc.create_delta_sync_index_and_wait(
        endpoint_name=VECTOR_SEARCH_ENDPOINT,  # 作成済みのエンドポイント
        index_name=CHUNKS_VECTOR_INDEX,  # ベクトル検索用インデックス名
        primary_key="chunk_id",  # 一意のキー（ドキュメントID）
        source_table_name=DELTA_TABLE_NAME,  # Delta Tableの名前
        pipeline_type="TRIGGERED",  # データが変わったら自動更新
        embedding_source_column="chunked_text",  # 埋め込むテキストの列名
        embedding_model_endpoint_name="databricks-bge-large-en",  # 埋め込み（ベクトル化）に使うモデル
    )
except Exception as e:
    print(f"Index {CHUNKS_VECTOR_INDEX} already exists. Skipping index creation.")

# COMMAND ----------

# DBTITLE 1,プロンプト
PROMPT_FORMAT_1 = """
        あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。
        会社の規定と照らし合わせて、申請が適切かどうかを判断してください。

        ### 重要な指示
        - 会社規定に基づいて判断し、独自の推論は行わないこと。
        - 判断結果は「問題なし」と「問題あり」の2つのみとする。
        - 規定に記載がない、または該当しない場合は、独自の推論をせず、必ず「問題なし」とみなすこと。
        - 数値データ（距離や金額など）は、そのままの単位（km、m、円）で扱い、変換しないこと。
        - 定期代が複数の経路に分かれている場合は、すべて合計した金額で判断すること。
        - 出力は必ず以下のフォーマットで返答すること。

        ### 出力フォーマット：
         [申請者名]さん
        - 問題なしの場合：
        「問題なし。通勤距離は適切です。」

        - 問題ありの場合：
        「問題あり。」
        「[引用した会社規定の内容][規程に基づいた詳細な判断理由]」
    """

PROMPT_FORMAT_2 = """  
    あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。
    会社の規定と照らし合わせて、申請が適切かどうかを判断してください。

    ### **重要な指示**
        - 会社規定に基づいて判断し、独自の推論は行わないこと。
        - 複数の経路に分かれている場合は、合計した金額が1ヶ月の定期代として判断すること。また、合計した距離が通勤距離として判断すること。
        - 判断結果は「問題なし」と「問題あり」の2パターンだけです。
        - 「問題なし」の場合、詳細な理由は記載せず、「通勤距離と定期代は適切です。」とだけ返答すること。
        - 「問題あり」の場合のみ、会社規定を引用し、その理由を明確に述べること。
        - 規定に記載がない、または該当しない場合は「問題なし」とみなすこと。
        - 出力は必ず以下のフォーマットで返答すること。

    ### **出力フォーマット**
    - **問題なしの場合**
      ```
      [申請者名]さん
      - 問題なし。通勤距離と定期代は適切です。
      ```

    - **問題ありの場合**
      ```
      [申請者名]さん
      - 問題あり。
      -[引用した会社規定の内容][規程に基づいた詳細な判断理由]
      ```
"""

PROMPT_FORMAT_NEW_1 = """  
    あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。
    会社の規定と照らし合わせて、申請が適切かどうかを判断してください。

    ### **重要な指示**
        - 会社規定に基づいて判断し、独自の推論は行わないこと。
        - 複数の経路に分かれている場合は、合計した金額が1ヶ月の定期代として判断すること。また、合計した距離が通勤距離として判断すること。
        - 判断結果は「問題なし」、「問題あり」、「要確認」の3パターンのみです。
        - 規定に記載がない場合、判断がつかない場合、例外がある場合は「要確認」と判断すること。
        - 出力は必ず以下のフォーマットで返答すること。

    ### **出力フォーマット**
      ```
      [申請者名]さん
      -[引用した会社規定の内容][規程に基づいた詳細な判断理由]
      -結果: 問題なし or 問題あり or 要確認
      ```
"""

PROMPT_FORMAT_NEW_2 = """  
    あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。
    会社の規定と照らし合わせて、申請が適切かどうかを判断してください。

    ### **重要な指示**
      - 会社規定に基づいて判断し、個人的な推論は行わないこと。
      - 経路が複数に分かれている場合は、合計金額を1か月の定期代として判断すること。また、合計距離を通勤距離として判断すること。
      - 最寄り駅に関しては【最寄り駅確認結果】を参考すること。
      - 判断結果は、「問題なし」「問題あり」「要確認」の3種類のみとする。
      - すべての規定に適合している場合は「問題なし」とし、一つでも適合していない場合は「問題あり」とすること。
      - 規定に記載がない場合、判断がつかない場合、または別途定める基準がある場合は、「要確認」とすること。
      - 必ずすべての申請者を漏れなくチェックすること。
      - 申請履歴に同じ申請者名の履歴があれば、申請履歴を表示すること。また、履歴がない場合は、申請履歴を表示しないこと。
      - 出力は必ず、以下のフォーマットで申請者毎に一人ずつ返答すること。

    ### **出力フォーマット**
      [申請者名]さん
      -過去の申請履歴: [申請履歴] or なし
      -[引用した会社規定の内容][規程に基づいた詳細な判断理由]
      -結果: 問題なし or 問題あり or 要確認
      -最寄り駅確認結果: 一致 or 不一致
"""

PROMPT_FORMAT_NEW_3 = """
  - あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。  
  - 会社の規定と照らし合わせて、申請が適切かどうかを判断してください。
  - 入力される会社規定や申請情報などはすべて**日本語**で記述されています。
  ---

  ### **重要な指示**

  - **会社規定に基づいて判断し、個人的な推論は行わないこと。**
  - 経路が複数に分かれている場合は、**合計金額を1か月の定期代**として判断すること。また、**合計距離を通勤距離**として判断すること。
  - **最寄り駅に関しては「【最寄り駅確認結果】」を必ず参考にすること。**
  - 判断結果は、**「問題なし」「問題あり」「要確認」** の3種類のみとする。
  - すべての規定に適合している場合は「問題なし」とし、**一つでも適合していない場合は「問題あり」** とすること。
  - **規定に記載がない場合、判断がつかない場合、または別途定める基準がある場合は「要確認」とすること。**
  - **すべての申請者を漏れなくチェックすること。**
  - 今回の申請と申請履歴に一致する申請者名があれば、履歴ありと判断し、その**申請履歴を表示すること**。履歴がない場合は、申請履歴のセクションは表示しないこと。

  ---

  ### **出力フォーマット（申請者ごとに1人ずつ）**

  #### [申請者名] さん

  - **過去の申請履歴**:  
    - [申請履歴内容]  
    - or `なし`

  - **会社規定の引用と判断理由**:  
    - `[引用した会社規定]`  
    - `[規程に基づいた詳細な判断理由]`

  - **結果**:  
    - `問題なし`  
    - `問題あり`  
    - `要確認`

  - **最寄り駅確認結果**:  
    - `一致`  
    - `不一致`
"""


# COMMAND ----------

# DBTITLE 1,RAGアプリの設定: LLM、ベクトル検索、プロンプトの構成
MODEL_NAME_1 = "databricks-dbrx-instruct"
MODEL_NAME_2 = "databricks-mixtral-8x7b-instruct"
MODEL_NAME_3 = "databricks-meta-llama-3-3-70b-instruct"
MODEL_NAME_4 = "databricks-meta-llama-3-1-8b-instruct"
MODEL_NAME_5 = "databricks-meta-llama-3-1-405b-instruct"

chain_config = { 
    "llm_model_serving_endpoint_name": MODEL_NAME_3, # LLMのエンドポイント
    "vector_search_endpoint_name": VECTOR_SEARCH_ENDPOINT,  # ベクトル検索エンドポイント
    "vector_search_index": CHUNKS_VECTOR_INDEX,  # 検索インデックス
    # プロンプト設定
    "llm_prompt_template": PROMPT_FORMAT_NEW_3,
}

# 入力データのサンプル
# input_example = {
#     "messages": [
#         {
#             "role": "user",
#             "content": """
#                 【通勤手当申請】
#                 申請者名: 田中太郎
#                 勤務先住所: 東京都渋谷区神宮前1-2-3 サンプル株式会社
#                 自宅住所: 東京都世田谷区成城8-9-10
#                 最寄り駅: 成城学園前駅
#                 利用交通機関と経路:
#                   - 成城学園前駅 → 下北沢駅（小田急線）
#                   - 下北沢駅 → 渋谷駅（京王井の頭線）
#                 通勤距離:
#                   - 小田急線: 8km
#                   - 京王井の頭線: 4km
#                 定期代:
#                   - 小田急線:
#                       - 1ヶ月: 9,000円
#                       - 3ヶ月: 25,500円
#                       - 6ヶ月: 48,000円
#                   - 京王井の頭線:
#                       - 1ヶ月: 6,000円
#                       - 3ヶ月: 16,500円
#                       - 6ヶ月: 32,000円
#                 ---
#                 申請者名: 鈴木花子
#                 勤務先住所: 神奈川県横浜市西区みなとみらい2-2-1 サンプル東京支社
#                 自宅住所: 東京都大田区蒲田5-6-7
#                 最寄り駅: 蒲田駅
#                 利用交通機関と経路:
#                   - 蒲田駅 → 横浜駅（京浜東北線）
#                   - 横浜駅 → みなとみらい駅（みなとみらい線）
#                 通勤距離:
#                   - 京浜東北線: 20km
#                   - みなとみらい線: 5km
#                 定期代:
#                   - 京浜東北線:
#                       - 1ヶ月: 14,000円
#                       - 3ヶ月: 39,000円
#                       - 6ヶ月: 75,000円
#                   - みなとみらい線:
#                       - 1ヶ月: 6,000円
#                       - 3ヶ月: 18,000円
#                       - 6ヶ月: 33,000円
#                 ---
#                 申請者名: 山田一郎
#                 勤務先住所: 千葉県千葉市美浜区中瀬2-6 サンプル千葉工場
#                 自宅住所: 千葉県船橋市本町7-8-9
#                 最寄り駅: 船橋駅
#                 利用交通機関と経路: 自家用車
#                 通勤距離: 18km
#                 定期代: なし
#                 ---
#                 申請者名: 佐藤美智子
#                 勤務先住所: 東京都港区六本木1-2-3 サンプル本社
#                 自宅住所: 埼玉県さいたま市大宮区大門町3-4-5
#                 最寄り駅: 大宮駅
#                 利用交通機関と経路:
#                   - 大宮駅 → 新宿駅（湘南新宿ライン）
#                   - 新宿駅 → 六本木駅（東京メトロ日比谷線）
#                 通勤距離:
#                   - 湘南新宿ライン: 25km
#                   - 東京メトロ日比谷線: 5km
#                 定期代:
#                   - 湘南新宿ライン:
#                       - 1ヶ月: 18,000円
#                       - 3ヶ月: 51,000円
#                       - 6ヶ月: 96,000円
#                   - 東京メトロ日比谷線:
#                       - 1ヶ月: 4,000円
#                       - 3ヶ月: 12,000円
#                       - 6ヶ月: 22,000円
#                 ---
#                 申請者名: 高橋太郎
#                 勤務先住所: 東京都中央区銀座5-6-7 サンプルビル
#                 自宅住所: 東京都台東区上野公園8-9-0
#                 最寄り駅: 上野駅
#                 利用交通機関と経路: 上野駅 → 銀座駅（東京メトロ銀座線）
#                 通勤距離: 8km
#                 定期代:
#                   - 1ヶ月: 12,000円
#                   - 3ヶ月: 34,000円
#                   - 6ヶ月: 65,000円
#                 ---
#                 申請者名: 木村春香
#                 勤務先住所: 神奈川県横浜市中区日本大通1-2-3 サンプル横浜支店
#                 自宅住所: 神奈川県川崎市高津区溝口1-2-3
#                 最寄り駅: 溝の口駅
#                 利用交通機関と経路: 溝の口駅 → 横浜駅（JR南武線、JR横須賀線） → 日本大通り駅（みなとみらい線）
#                 通勤距離: 16km
#                 定期代:
#                   - 1ヶ月: 18,000円
#                   - 3ヶ月: 52,000円
#                   - 6ヶ月: 100,000円
#                 ---
#                 申請者名: 山口優子
#                 勤務先住所: 千葉県浦安市美浜1-2-3 サンプル東京第二工場
#                 自宅住所: 東京都江戸川区東小岩7-8-9
#                 最寄り駅: 小岩駅
#                 利用交通機関と経路: 小岩駅 → 新浦安駅（JR総武線、JR京葉線）
#                 通勤距離: 40km
#                 定期代:
#                   - 1ヶ月: 45,000円
#                   - 3ヶ月: 130,000円
#                   - 6ヶ月: 250,000円
#                 ---
#                 申請者名: 加藤美咲
#                 勤務先住所: 東京都渋谷区神南1-2-3 サンプルグループ本社
#                 自宅住所: 東京都世田谷区用賀4-5-6
#                 最寄り駅: 用賀駅
#                 利用交通機関と経路: 用賀駅 → 渋谷駅（東急田園都市線）
#                 通勤距離: 10km
#                 定期代:
#                   - 1ヶ月: 14,000円
#                   - 3ヶ月: 40,000円
#                   - 6ヶ月: 77,000円
#                 ---
#                 申請者名: 佐々木翔
#                 勤務先住所: 東京都港区赤坂1-3-4 サンプル赤坂オフィス
#                 自宅住所: 東京都港区赤坂5-6-7
#                 最寄り駅: 赤坂駅
#                 利用交通機関と経路: 赤坂駅 → 赤坂見附駅（東京メトロ千代田線）
#                 通勤距離: 1.2km（バス利用）
#                 定期代:
#                   - バス:
#                       - 1ヶ月: 5,000円
#                       - 3ヶ月: 14,500円
#                       - 6ヶ月: 28,000円
#                 ---
#                 申請者名: 大谷健
#                 勤務先住所: 東京都品川区大崎2-3-4 サンプル大崎支社
#                 自宅住所: 東京都立川市柴崎町1-2-3
#                 最寄り駅: 立川駅
#                 利用交通機関と経路:
#                   - 立川駅 → 東京駅（JR中央線）
#                   - 東京駅 → 大崎駅（JR山手線）
#                 通勤距離:
#                   - JR中央線: 40km
#                   - JR山手線: 5km
#                 定期代:
#                   - JR中央線:
#                       - 1ヶ月: 28,000円
#                       - 3ヶ月: 81,000円
#                       - 6ヶ月: 155,000円
#                   - JR山手線:
#                       - 1ヶ月: 10,000円
#                       - 3ヶ月: 29,000円
#                       - 6ヶ月: 55,000円
#             """
#         }
#     ]
# }

input_example = { 
    "messages": [
        {
            "role": "user",
            "content": """
                以下はいろいろな形式の申請書です。
                すべての申請内容をチェックしてください。
                """
        }
    ]
}

# COMMAND ----------

# DBTITLE 1,マルチタスク　テスト
import os

# Databricksのワークスペースフォルダのパス
folder_path = "multitasking/"

# フォルダ内のファイルをリストアップ
files = [os.path.join(folder_path, f) for f in os.listdir(folder_path)]

all_text = ""

for file in files:
    try:
        extracted_text = extract_text(file)
        all_text += f"\n### {os.path.basename(file)} の内容 ###\n{extracted_text}\n"
    except Exception as e:
        print(f"エラー: {file} の処理に失敗しました: {e}")

print(all_text)


# COMMAND ----------

# DBTITLE 1,MLflowにRAGチェーンを登録し、ローカルで動作確認
import mlflow

# MLflowの実験（Run）を開始する（run_name は "databricks-docs-bot"）
with mlflow.start_run(run_name="databricks-docs-bot"):
    
    # MLflowにRAGチェーン（LangChainのモデル）を保存する
    logged_chain_info = mlflow.langchain.log_model(
        # lc_model: RAGチェーンのスクリプトファイルのパス
        lc_model=os.path.join(os.getcwd(), "rag_chain"),
        # model_config: RAG チェーンの設定情報（どの LLM を使うか、プロンプトのテンプレートなど）
        model_config=chain_config,
        # artifact_path: MLflow に保存する際の保存先フォルダ（"chain" という名前で保存）
        artifact_path="chain",  
        # input_example: RAG チェーンに入力するデータのフォーマット例（どんな形の入力を受け取るか）
        input_example=input_example, 
        # pip_requirements: RAG チェーンのスクリプトファイルに必要なpip パッケージ
        pip_requirements=["langchain", "googlemaps"],
    )

# Test the chain locally to see the MLflow Trace
chain = mlflow.langchain.load_model(logged_chain_info.model_uri)
chain.invoke(input_example)

# COMMAND ----------

# DBTITLE 1,Agent Frameworkを使って、RAGアプリケーションをデプロイ
from databricks import agents  # RAGアプリをデプロイするために Databricksのagents API を使う
import time  # デプロイ完了を待機するために使う
from databricks.sdk.service.serving import EndpointStateReady, EndpointStateConfigUpdate  # エンドポイントの状態をチェックするために使う

# Use Unity Catalog to log the chain
mlflow.set_registry_uri('databricks-uc')


# Register the chain to UC
uc_registered_model_info = mlflow.register_model(model_uri=logged_chain_info.model_uri, name=UC_MODEL_NAME)

# Deploy to enable the Review APP and create an API endpoint
deployment_info = agents.deploy(model_name=UC_MODEL_NAME, model_version=uc_registered_model_info.version)

# Wait for the Review App to be ready
print("\nWaiting for endpoint to deploy.  This can take 15 - 20 minutes.", end="")
while w.serving_endpoints.get(deployment_info.endpoint_name).state.ready == EndpointStateReady.NOT_READY or w.serving_endpoints.get(deployment_info.endpoint_name).state.config_update == EndpointStateConfigUpdate.IN_PROGRESS:
    print(".", end="")
    time.sleep(30)

print(f"Endpoint {deployment_info.endpoint_name} is now ready!")
