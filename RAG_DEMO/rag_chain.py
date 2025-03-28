# Databricks notebook source
from operator import itemgetter
import mlflow
import os
from pypdf import PdfReader
from docx import Document
import pandas as pd
from pyspark.sql import SparkSession
from typing import Optional

from databricks.vector_search.client import VectorSearchClient

from langchain_community.chat_models import ChatDatabricks
from langchain_community.vectorstores import DatabricksVectorSearch
from langchain_core.runnables import RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import (
    PromptTemplate,
    ChatPromptTemplate,
    MessagesPlaceholder,
    SystemMessagePromptTemplate,
)
from langchain_core.runnables import RunnablePassthrough, RunnableBranch
from langchain_core.messages import HumanMessage, AIMessage
from langchain.agents import create_tool_calling_agent, AgentExecutor
from langchain_core.tools import tool
from typing import List
import googlemaps

## Enable MLflow Tracing
mlflow.langchain.autolog()

############
# Helper functions
############
# Return the string contents of the most recent message from the user
def extract_user_query_string(chat_messages_array):
    return chat_messages_array[-1]["content"]

# Return the chat history, which is is everything before the last question
def extract_chat_history(chat_messages_array):
    return chat_messages_array[:-1]

def refine_query(query):
  return (
    "以下の点に注意して、会社規定を確認してください：\n\n"
    "- **支給対象者**\n"
    "- **支給経路の決定**\n"
    "- **支給額**\n\n"
    f"## \n\n{query.strip()}"
  )


def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""

    # 通常のテキスト（段落）
    for para in doc.paragraphs:
        if para.text.strip():  # 空白行を削除
            text += para.text.strip() + "\n"

    # テーブルデータの取得
    for table in doc.tables:
        processed_cells = set()  # 既に処理したセルを保存

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if (cell._tc not in processed_cells) and cell_text:  # 結合セルを除外
                    row_data.append(cell_text)
                    processed_cells.add(cell._tc)
                else:
                    row_data.append("")

            text += " | ".join(row_data) + "\n"

    return text

def extract_text_from_xlsx(xlsx_path):
    df = pd.read_excel(xlsx_path, engine='openpyxl', sheet_name=None) 
    text = ""
    for sheet_name, sheet_df in df.items():
        text += f"### {sheet_name} シート ###\n"
        text += sheet_df.to_string(index=False, header=True) + "\n"
    return text

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="shift_jis") as f:
        text = f.read()
    return text

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".xls", ".xlsx"]:
        return extract_text_from_xlsx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"対応していないファイル形式です: {ext}")

@tool
def find_nearest_stations(addresses: List[str]) -> List[str]:
  """複数の住所から、それぞれの最寄り駅を返します"""
  api_key = "AIzaSyDU8HiLkRbwzqG2kCljLiUvD1TyuTi6JpQ"
  gmaps = googlemaps.Client(key=api_key)

  results = []
  for address in addresses:
    geocode_result = gmaps.geocode(address)
    if not geocode_result:
      results.append(f"{address} → 住所が見つかりませんでした")
      continue

    location = geocode_result[0]['geometry']['location']
    lat, lng = location['lat'], location['lng']

    places = gmaps.places_nearby(location=(lat, lng), radius=1000, type='train_station')
    if places.get("results"):
      station_name = places['results'][0]['name']
      results.append(f"{address} → 最寄り駅: {station_name}")
    else:
      results.append(f"{address} → 最寄り駅が見つかりませんでした")

  return results

# Load the chain's configuration
model_config = mlflow.models.ModelConfig(development_config="rag_chain_config.yaml")

############
# Connect to the Vector Search Index
############
vs_client = VectorSearchClient(disable_notice=True)
vs_index = vs_client.get_index(
    endpoint_name=model_config.get("vector_search_endpoint_name"),
    index_name=model_config.get("vector_search_index"),
)

############
# Turn the Vector Search index into a LangChain retriever
############
vector_search_as_retriever = DatabricksVectorSearch(
    vs_index,
    text_column="chunked_text",
    columns=[
        "chunk_id",
        "chunked_text",
    ],
).as_retriever(search_kwargs={"k":18, "score_threshold": 0.5})

############
# Required to:
# 1. Enable the RAG Studio Review App to properly display retrieved chunks
# 2. Enable evaluation suite to measure the retriever
############

mlflow.models.set_retriever_schema(
    primary_key="chunk_id",
    text_column="chunked_text",
    doc_uri="url",  # Review App uses `doc_uri` to display chunks from the same document in a single view
)


############
# Method to format the docs returned by the retriever into the prompt
############
def format_context(docs):
  chunk_template = (
    "{chunk_text}\n\n"
  )
  chunk_contents = [
    chunk_template.format(
      index=i + 1,
      chunk_text=d.page_content.strip()
    )
    for i, d in enumerate(docs)
  ]

  contents = "".join(chunk_contents)

  return f"{contents.strip()}\n\n"

# SparkSession を取得
spark = SparkSession.builder.getOrCreate()

# データベースのテーブル情報
DB_NAME = "hhhd_demo_itec.allowance_payment_rules"
TABLE_NAME = "commuting_allowance_history"

# カラム名のマッピング
COLUMN_MAPPING = {
    "name": "申請者名",
    "work_address": "勤務先住所",
    "address": "自宅住所",
    "nearest_station": "最寄り駅",
    "route_1": "利用交通機関と経路①",
    "distance_1": "通勤距離①",
    "commuter_pass_1": "定期代①",
    "route_2": "利用交通機関と経路②",
    "distance_2": "通勤距離②",
    "commuter_pass_2": "定期代②"
}

def format_history_text(records):
  if not records:
    return "履歴なし"

  formatted_history = []
  for i, row in enumerate(records):
    formatted_row = {
      COLUMN_MAPPING.get(col, col): row[col] for col in row.asDict()
    }
    formatted_history.append(
      f"- **申請者名**: {formatted_row['申請者名']}\n"
      f"- **勤務先住所**: {formatted_row['勤務先住所']}\n"
      f"- **自宅住所**: {formatted_row['自宅住所']}\n"
      f"- **最寄り駅**: {formatted_row['最寄り駅']}\n"
      f"- **利用交通機関と経路①**: {formatted_row['利用交通機関と経路①']}\n"
      f"- **通勤距離①**: {formatted_row['通勤距離①']} km\n"
      f"- **定期代①**: {formatted_row['定期代①']} 円\n"
      f"- **利用交通機関と経路②**: {formatted_row['利用交通機関と経路②']}\n"
      f"- **通勤距離②**: {formatted_row['通勤距離②']} km\n"
      f"- **定期代②**: {formatted_row['定期代②']} 円\n"
    )

  history_text = "\n---\n\n".join(formatted_history)
  return f"\n\n{history_text}"


# 申請履歴を取得
def get_user_history_from_db(data) -> str:
    query = f"""
    SELECT name, work_address, address, nearest_station, 
           route_1, distance_1, commuter_pass_1, 
           route_2, distance_2, commuter_pass_2
    FROM hhhd_demo_itec.allowance_payment_rules.commuting_allowance_history
    """

    df = spark.sql(query)
    records = df.collect()

    return format_history_text(records)

############
# Prompt Template for generation
############
def get_llm_prompt_template(data):
    return model_config.get("llm_prompt_template")

system_template = SystemMessagePromptTemplate.from_template(
    """
    {llm_prompt_template}

    ---

    ## 会社規定
    {context}

    ---

    ## 申請履歴
    {user_history}

    ---

    ## 最寄り駅確認結果
    {address_check_result}
    """
)

prompt = ChatPromptTemplate.from_messages([
  system_template,
  MessagesPlaceholder(variable_name="formatted_chat_history"),
  ("user", "{question}"),
])

# Format the converastion history to fit into the prompt template above.
def format_chat_history_for_prompt(chat_messages_array):
    history = extract_chat_history(chat_messages_array)
    formatted_chat_history = []
    if len(history) > 0:
        for chat_message in history:
            if chat_message["role"] == "user":
                formatted_chat_history.append(
                    HumanMessage(content=chat_message["content"])
                )
            elif chat_message["role"] == "assistant":
                formatted_chat_history.append(
                    AIMessage(content=chat_message["content"])
                )
    return formatted_chat_history


############
# Prompt Template for query rewriting to allow converastion history to work - this will translate a query such as "how does it work?" after a question such as "what is spark?" to "how does spark work?".
############
query_rewrite_template = """Based on the chat history below, we want you to generate a query for an external data source to retrieve relevant documents so that we can better answer the question. The query should be in natural language. The external data source uses similarity search to search for relevant documents in a vector space. So the query should be similar to the relevant documents semantically. Answer with only the query. Do not add explanation.

Chat history: {chat_history}

Question: {question}"""

query_rewrite_prompt = PromptTemplate(
    template=query_rewrite_template,
    input_variables=["chat_history", "question"],
)


############
# FM for generation
############
model = ChatDatabricks(
    endpoint=model_config.get("llm_model_serving_endpoint_name"),
    extra_params={"temperature": 0.01},
)

############
# questionから情報を抽出して、定型文に変換するChain　
############
extract_prompt_template = PromptTemplate.from_template("""
    申請書データ:
    {input_text}                                                      

    上記の申請書データから、以下の情報を抽出してください：

    - 申請者名
    - 勤務先住所
    - 自宅住所
    - 最寄り駅
    - 利用交通機関と経路
    - 通勤距離
    - 定期代

    複数の交通機関と経路がある場合は、それぞれについて番号を付けて記載してください。
    複数申請者がある場合は、それぞれについて番号を付けて記載してください。
    出力は以下のMarkdown形式に従ってください：

    ---

    ### 申請者情報
    - **申請者①**
    - **申請者名**:  
    - **勤務先住所**:  
    - **自宅住所**:  
    - **最寄り駅**:  

    ### 通勤経路①

    - **利用交通機関と経路①**:  
    - **通勤距離①**:  
    - **定期代①**:  

    ### 通勤経路②（ある場合のみ）

    - **利用交通機関と経路②**:  
    - **通勤距離②**:  
    - **定期代②**:

    ---
    """)

extract_fields_chain = extract_prompt_template | model | StrOutputParser()

def extract_structured_fields_with_llm(text: str) -> str:
    return extract_fields_chain.invoke({"input_text": text})

def extract_text_from_file(query):
    # Databricksのワークスペースフォルダのパス
    # folder_path = "multitasking/"
    folder_path = "test_from/"

    files = [os.path.join(folder_path, f) for f in os.listdir(folder_path)]

    all_text = ""

    for file in files:
        try:
            extracted_text = extract_text(file)
            all_text += f"\n---\n\n### {os.path.basename(file)} の申請内容\n\n```\n{extracted_text}\n```\n"
        except Exception as e:
            print(f"エラー: {file} の処理に失敗しました: {e}")

    llm_text = extract_structured_fields_with_llm(all_text)

    return (
            f"## \n\n{query.strip()}"
            f"\n\n{llm_text}"
        )
    
############
# 駅検索するために住所を英語表示に変換するChain　
############
change_address_prompt_template = PromptTemplate.from_template("""\
申請書データ：
{input_text}

---

上記の申請書データから、以下の情報を申請者ごとに抽出してください：

- 申請者名
- 自宅住所
- 最寄り駅

複数の申請者が含まれる場合は、それぞれ `申請者①`, `申請者②` のように番号を付けて記載してください。

出力は以下のMarkdown形式に厳密に従ってください：

---

### 申請者情報

#### 申請者①
- **申請者名**:  
- **自宅住所**:  
- **最寄り駅**:  

#### 申請者②
- **申請者名**:  
- **自宅住所**:  
- **最寄り駅**:  

...（以下続く）
""")

change_address_chain = change_address_prompt_template | model | StrOutputParser()

def change_address_with_llm(text: str) -> str:
    return change_address_chain.invoke({"input_text": text})

################################################
# 最寄り駅をチェックするagent
################################################
# ツール一覧
tools = [find_nearest_stations]

station_check_instruction = """
    あなたは通勤手当申請のチェックを行うアシスタントです。

    申請書データから、最寄り駅を検索してください。
    その上で、「申請書に記載された最寄り駅」と一致するかどうかを確認し、結果を出力してください。

    - 出力フォーマットは以下の通りです（この形式に厳密に従ってください）：

    #### 申請者①
    - 申請者名: 
    - 申請書に記載された最寄り駅:
    - 検索された最寄り駅:
    - 判断結果:

    #### 申請者②
    - 申請者名: 
    - 申請書に記載された最寄り駅:
    - 検索された最寄り駅:
    - 判断結果:

    ...（以下続く）
    """

agent_prompt = ChatPromptTemplate.from_messages([
  ("system", station_check_instruction),
  ("user", "{input_text}"),
  ("placeholder", "{agent_scratchpad}")
])

agent = create_tool_calling_agent(model, tools, agent_prompt)
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

# check_station_match_with_llm関数
def check_station_match_with_llm(text: str) -> str:
  input_text = f"申請書データ: {text}"
  # result = agent_executor.invoke({"input_text": input_text})
  result = agent_executor.invoke({
    "input_text": input_text,
    "agent_scratchpad": []
  })
  return result["output"] if isinstance(result, dict) and "output" in result else str(result)

def check_station_match(text: str) -> str:
    changed_address_text = change_address_with_llm(text)
    return check_station_match_with_llm(changed_address_text)

############
# Main RAG Chain
############
chain = (
    {
        "question": itemgetter("messages") | RunnableLambda(extract_user_query_string) | RunnableLambda(refine_query) | RunnableLambda(extract_text_from_file),
        "chat_history": itemgetter("messages") | RunnableLambda(extract_chat_history),
        "formatted_chat_history": itemgetter("messages") | RunnableLambda(format_chat_history_for_prompt),
        "user_history": itemgetter("messages") | RunnableLambda(get_user_history_from_db),
    }
    | RunnablePassthrough()
    | {
        "context": RunnableBranch(  # Only re-write the question if there is a chat history
            (
                lambda x: len(x["chat_history"]) > 0,
                query_rewrite_prompt | model | StrOutputParser(),
            ),
            itemgetter("question"),
        )
        | vector_search_as_retriever
        | RunnableLambda(format_context),
        "formatted_chat_history": itemgetter("formatted_chat_history"),
        "question": itemgetter("question"),
        "user_history": itemgetter("user_history"),
        "address_check_result": itemgetter("question") | RunnableLambda(check_station_match),
        "llm_prompt_template": RunnableLambda(get_llm_prompt_template),
    }
    | prompt
    | model
    | StrOutputParser()
)

## Tell MLflow logging where to find your chain.
# `mlflow.models.set_model(model=...)` function specifies the LangChain chain to use for evaluation and deployment.  This is required to log this chain to MLflow with `mlflow.langchain.log_model(...)`.

mlflow.models.set_model(model=chain)
